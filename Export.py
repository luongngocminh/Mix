# -*- coding: UTF-8 -*-
from docx import Document

answerIndex = ['A. ','B. ','C. ','D. ']
newDoc = Document()
def export(questions):
	for i in range(len(questions)):
		s = ''
		answs = questions[i].answers
		paragraph = newDoc.add_paragraph()
		paragraph.add_run(u'Câu {0}: '.format(str(i+1))).bold = True
		paragraph.add_run(questions[i].question)
		paragraph = newDoc.add_paragraph()
		for j in range(len(answs)):
			paragraph.add_run(answerIndex[j]).bold = True
			paragraph.add_run(answs[j])
			paragraph.add_run('        ')

	paragraph = newDoc.add_paragraph('Answers: ')
	for i in range(len(questions)):
		paragraph = newDoc.add_paragraph(str(i+1)+u'. '+questions[i].rightAnswerIndex)
	newDoc.save('Documents/new_Test.docx')

